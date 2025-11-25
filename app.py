import streamlit as st
from docx import Document
from io import BytesIO

st.set_page_config(page_title="Gerador de Matriz de Riscos", layout="centered")

st.title("🏛️ Gerador de Matriz de Riscos")
st.markdown("---")
st.info("Ferramenta de apoio às Unidades de Controle Interno Municipal.")

# Formulário
with st.form("form_auditoria"):
    st.subheader("1. Mapeamento")
    processo = st.text_input("Nome do Processo", placeholder="Ex: Licitação - Pregão Eletrônico")
    responsavel = st.text_input("Responsável pelo Processo")
    
    st.subheader("2. Análise de Risco")
    risco_desc = st.text_area("Descrição do Risco Identificado", placeholder="O que pode dar errado?")
    col1, col2 = st.columns(2)
    with col1:
        prob = st.slider("Probabilidade (1 = Raro, 5 = Quase Certo)", 1, 5)
    with col2:
        impacto = st.slider("Impacto (1 = Insignificante, 5 = Catastrófico)", 1, 5)
        
    nivel_risco = prob * impacto
    
    submitted = st.form_submit_button("Gerar Documento Word")

# Ação após clicar no botão
if submitted:
    # Cria o documento na memória
    document = Document()
    
    document.add_heading('Matriz de Riscos e Controles', 0)
    document.add_paragraph(f'Processo Auditado: {processo}')
    document.add_paragraph(f'Responsável: {responsavel}')
    document.add_paragraph('---')
    
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Risco'
    hdr_cells[1].text = 'Nível (PxI)'
    hdr_cells[2].text = 'Classificação'
    
    row_cells = table.add_row().cells
    row_cells[0].text = risco_desc
    row_cells[1].text = str(nivel_risco)
    
    if nivel_risco >= 15:
        row_cells[2].text = "CRÍTICO"
    elif nivel_risco >= 8:
        row_cells[2].text = "MÉDIO"
    else:
        row_cells[2].text = "BAIXO"

    # Prepara o download
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    st.success(f"Documento gerado! Risco calculado: {nivel_risco}")
    st.download_button(
        label="📥 Baixar Matriz (.docx)",
        data=buffer,
        file_name=f"Matriz_Riscos_{processo}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
